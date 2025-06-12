import io
import re
from typing import Dict, Optional
from pathlib import Path

import PyPDF2
from docx import Document
import pandas as pd

from ..models.evaluation import DocumentType, ParsedDocument


class DocumentParser:
    """Utility class for parsing various document formats."""
    
    @staticmethod
    def parse_document(file_content: bytes, filename: str) -> ParsedDocument:
        """Parse document based on file extension."""
        file_path = Path(filename)
        extension = file_path.suffix.lower()
        
        if extension == '.pdf':
            return DocumentParser._parse_pdf(file_content, filename)
        elif extension == '.docx':
            return DocumentParser._parse_docx(file_content, filename)
        elif extension in ['.txt', '.md']:
            return DocumentParser._parse_text(file_content, filename)
        elif extension in ['.xlsx', '.xls']:
            return DocumentParser._parse_excel(file_content, filename)
        else:
            raise ValueError(f"Unsupported file format: {extension}")
    
    @staticmethod
    def _parse_pdf(file_content: bytes, filename: str) -> ParsedDocument:
        """Parse PDF document."""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            content = ""
            for page in pdf_reader.pages:
                content += page.extract_text() + "\n"
            
            sections = DocumentParser._extract_sections(content)
            document_type = DocumentParser._detect_document_type(content)
            
            return ParsedDocument(
                content=content,
                document_type=document_type,
                sections=sections,
                metadata={
                    "filename": filename,
                    "pages": str(len(pdf_reader.pages)),
                    "format": "pdf"
                }
            )
        except Exception as e:
            raise ValueError(f"Error parsing PDF: {str(e)}")
    
    @staticmethod
    def _parse_docx(file_content: bytes, filename: str) -> ParsedDocument:
        """Parse DOCX document."""
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            content = ""
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            # Also extract from tables
            table_content = ""
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            table_content += cell_text + " "
                    table_content += "\n"
            
            if table_content:
                content += "\n" + table_content
            
            sections = DocumentParser._extract_sections(content)
            document_type = DocumentParser._detect_document_type(content)
            
            return ParsedDocument(
                content=content,
                document_type=document_type,
                sections=sections,
                metadata={
                    "filename": filename,
                    "paragraphs": str(len(doc.paragraphs)),
                    "tables": str(len(doc.tables)),
                    "format": "docx"
                }
            )
        except Exception as e:
            raise ValueError(f"Error parsing DOCX: {str(e)}")
    
    @staticmethod
    def _parse_text(file_content: bytes, filename: str) -> ParsedDocument:
        """Parse text document."""
        try:
            content = file_content.decode('utf-8')
            sections = DocumentParser._extract_sections(content)
            document_type = DocumentParser._detect_document_type(content)
            
            return ParsedDocument(
                content=content,
                document_type=document_type,
                sections=sections,
                metadata={
                    "filename": filename,
                    "lines": str(len(content.split('\n'))),
                    "format": "text"
                }
            )
        except Exception as e:
            raise ValueError(f"Error parsing text file: {str(e)}")
    
    @staticmethod
    def _parse_excel(file_content: bytes, filename: str) -> ParsedDocument:
        """Parse Excel document."""
        try:
            excel_file = io.BytesIO(file_content)
            
            # Determine the appropriate engine based on file extension
            file_path = Path(filename)
            extension = file_path.suffix.lower()
            
            if extension == '.xlsx':
                engine = 'openpyxl'
            elif extension == '.xls':
                engine = 'xlrd'
            else:
                # Default to openpyxl for unknown extensions
                engine = 'openpyxl'
            
            # Read all sheets from the Excel file
            excel_data = pd.read_excel(excel_file, sheet_name=None, engine=engine)
            
            content = ""
            sheet_info = {}
            
            for sheet_name, df in excel_data.items():
                # Convert DataFrame to text representation
                sheet_content = f"\n=== Sheet: {sheet_name} ===\n"
                
                # Add column headers
                if not df.empty:
                    headers = " | ".join(str(col) for col in df.columns)
                    sheet_content += f"Columns: {headers}\n\n"
                    
                    # Add data rows (limit to first 100 rows to avoid overwhelming content)
                    for idx, row in df.head(100).iterrows():
                        row_data = " | ".join(str(val) if pd.notna(val) else "" for val in row.values)
                        sheet_content += f"{row_data}\n"
                    
                    if len(df) > 100:
                        sheet_content += f"\n... and {len(df) - 100} more rows\n"
                else:
                    sheet_content += "Empty sheet\n"
                
                content += sheet_content
                sheet_info[sheet_name] = {
                    "rows": len(df),
                    "columns": len(df.columns) if not df.empty else 0,
                    "column_names": list(df.columns) if not df.empty else []
                }
            
            sections = DocumentParser._extract_excel_sections(excel_data)
            document_type = DocumentType.OTHER  # Excel files are typically data/discovery documents
            
            return ParsedDocument(
                content=content,
                document_type=document_type,
                sections=sections,
                metadata={
                    "filename": filename,
                    "format": "excel",
                    "engine": engine,
                    "sheets": list(excel_data.keys()),
                    "sheet_info": sheet_info,
                    "total_sheets": len(excel_data)
                }
            )
        except Exception as e:
            raise ValueError(f"Error parsing Excel file: {str(e)}")
    
    @staticmethod
    def _extract_excel_sections(excel_data: Dict[str, pd.DataFrame]) -> Dict[str, str]:
        """Extract sections from Excel data based on sheet names and content."""
        sections = {}
        
        for sheet_name, df in excel_data.items():
            if not df.empty:
                # Create a summary of the sheet content
                summary = f"Sheet '{sheet_name}' contains {len(df)} rows and {len(df.columns)} columns.\n"
                
                if len(df.columns) > 0:
                    summary += f"Columns: {', '.join(str(col) for col in df.columns)}\n"
                
                # Add sample data if available
                if len(df) > 0:
                    summary += "\nSample data:\n"
                    sample_rows = min(3, len(df))
                    for idx, row in df.head(sample_rows).iterrows():
                        row_data = " | ".join(str(val) if pd.notna(val) else "N/A" for val in row.values)
                        summary += f"  {row_data}\n"
                
                sections[sheet_name.lower().replace(' ', '_')] = summary
        
        return sections
    
    @staticmethod
    def _extract_sections(content: str) -> Dict[str, str]:
        """Extract sections from document content."""
        sections = {}
        
        # Common section patterns
        section_patterns = [
            r'(?i)^(executive summary|summary)[\s:]*\n(.*?)(?=\n\s*[A-Z][^a-z]*\n|\Z)',
            r'(?i)^(introduction|overview)[\s:]*\n(.*?)(?=\n\s*[A-Z][^a-z]*\n|\Z)',
            r'(?i)^(approach|methodology)[\s:]*\n(.*?)(?=\n\s*[A-Z][^a-z]*\n|\Z)',
            r'(?i)^(solution|proposed solution)[\s:]*\n(.*?)(?=\n\s*[A-Z][^a-z]*\n|\Z)',
            r'(?i)^(timeline|schedule|project timeline)[\s:]*\n(.*?)(?=\n\s*[A-Z][^a-z]*\n|\Z)',
            r'(?i)^(team|resources|staffing)[\s:]*\n(.*?)(?=\n\s*[A-Z][^a-z]*\n|\Z)',
            r'(?i)^(technology|technical approach)[\s:]*\n(.*?)(?=\n\s*[A-Z][^a-z]*\n|\Z)',
            r'(?i)^(migration|migration approach)[\s:]*\n(.*?)(?=\n\s*[A-Z][^a-z]*\n|\Z)',
            r'(?i)^(assessment|analysis)[\s:]*\n(.*?)(?=\n\s*[A-Z][^a-z]*\n|\Z)',
            r'(?i)^(operations|operational support)[\s:]*\n(.*?)(?=\n\s*[A-Z][^a-z]*\n|\Z)',
        ]
        
        for pattern in section_patterns:
            matches = re.finditer(pattern, content, re.MULTILINE | re.DOTALL)
            for match in matches:
                section_name = match.group(1).lower().strip()
                section_content = match.group(2).strip()
                if section_content:
                    sections[section_name] = section_content
        
        return sections
    
    @staticmethod
    def _detect_document_type(content: str) -> DocumentType:
        """Detect document type based on content."""
        content_lower = content.lower()
        
        # RFP indicators
        rfp_indicators = [
            'request for proposal', 'rfp', 'request for quotation', 'rfq',
            'invitation to tender', 'itt', 'statement of work', 'sow'
        ]
        
        # Proposal indicators
        proposal_indicators = [
            'proposal', 'response to rfp', 'technical proposal',
            'commercial proposal', 'bid response'
        ]
        
        # Response indicators
        response_indicators = [
            'response', 'reply', 'submission', 'tender response'
        ]
        
        if any(indicator in content_lower for indicator in rfp_indicators):
            return DocumentType.RFP
        elif any(indicator in content_lower for indicator in proposal_indicators):
            return DocumentType.PROPOSAL
        elif any(indicator in content_lower for indicator in response_indicators):
            return DocumentType.RESPONSE
        else:
            return DocumentType.OTHER 